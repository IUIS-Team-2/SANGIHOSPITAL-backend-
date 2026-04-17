import os
try:
    import docx
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
except ImportError:
    print("❌ Error: Missing libraries. Run: pip install python-docx pandas openpyxl")
    exit()

def convert_word_to_excel(word_path, excel_path):
    print(f"📂 Reading Word Document: {word_path}...")
    
    if not os.path.exists(word_path):
        print(f"❌ Error: Could not find '{word_path}'. Make sure it's in the same folder.")
        return

    doc = docx.Document(word_path)
    all_data = []
    
    print(f"🔍 Found {len(doc.tables)} tables. Extracting data row by row...")
    for table in doc.tables:
        for row in table.rows:
            # Extract text from cells, replacing newlines with spaces
            row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            if any(row_data):
                all_data.append(row_data)
                
    # Initialize our perfect Excel Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Cash Rates"
    
    # We name the standard Non-NABH rate column "Rate" so import_data.py automatically picks it up!
    headers = ["Category", "Description", "Rate (NABH)", "Rate"]
    ws.append(headers)
    
    # Professional Styling
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), 
                         top=Side(style='thin'), bottom=Side(style='thin'))
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin_border
        
    current_category = "GENERAL SERVICES"
    count = 0
    
    for row in all_data:
        # Deduplicate cells (Word tables sometimes cause duplicate adjacent cell text)
        clean_row = []
        for item in row:
            if not clean_row or clean_row[-1] != item:
                clean_row.append(item)
                
        clean_row = [c for c in clean_row if c]
        
        # Detect Category Headers
        if len(clean_row) == 1:
            text = clean_row[0].upper()
            if len(text) > 3 and "RATE" not in text and "INDEX" not in text:
                current_category = text
            continue
            
        if not clean_row: continue
        
        # Remove "Sr. No" if it exists at the start of the row
        if clean_row[0].isdigit() or clean_row[0].replace('.','',1).isdigit():
            clean_row = clean_row[1:]
            
        if not clean_row: continue
        
        description = clean_row[0]
        if "description" in description.lower() or "procedure" in description.lower():
            continue 
            
        nabh = ""
        non_nabh = ""
        
        if len(clean_row) >= 3:
            nabh = clean_row[1]
            non_nabh = clean_row[2]
        elif len(clean_row) == 2:
            non_nabh = clean_row[1]
            
        def clean_price(val):
            v = val.replace('Rs.', '').replace(',', '').strip()
            try: return float(v)
            except: return val

        nabh_val = clean_price(nabh)
        non_nabh_val = clean_price(non_nabh)
        
        # Skip rows that don't contain real numbers for prices
        if isinstance(non_nabh_val, str) and not any(c.isdigit() for c in non_nabh_val):
            continue
            
        ws.append([current_category, description, nabh_val, non_nabh_val])
        count += 1
        
    # Apply borders and number formatting
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=4):
        for cell in row:
            cell.border = thin_border
            if cell.column in [3, 4] and isinstance(cell.value, (int, float)):
                cell.number_format = '#,##0.00'
                
    # Adjust widths to make it highly readable
    ws.column_dimensions['A'].width = 25
    ws.column_dimensions['B'].width = 60
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 15
    
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions

    wb.save(excel_path)
    print(f"\n🎉 SUCCESS! Extracted {count} services perfectly.")
    print(f"📁 The file has been saved to: {excel_path}")

if __name__ == "__main__":
    # Ensure the data folder exists
    if not os.path.exists('./data'):
        os.makedirs('./data')
        
    # Put the Word document in the same folder as this script!
    convert_word_to_excel("NEW RATE LIST SANGI HOSPITAL.docx", "./data/CASH_PRICES.xlsx")